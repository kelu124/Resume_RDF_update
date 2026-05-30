"""
resume_rdf.parsing
==================
Helpers for building Anthropic API content blocks from CV files and
processing the Turtle RDF output returned by the model.

Supported input formats
-----------------------
- ``.pdf``   — sent as a base64 document block (native Claude vision)
- ``.docx``  — text extracted via python-docx, sent as a text block
- ``.txt`` / ``.md`` / anything else — read as UTF-8 text block

python-docx is an optional dependency required only for ``.docx`` inputs::

    pip install python-docx
    # or
    pip install "resume-rdf[docx]"
"""

import base64
import io
import os
import re


# ── docx helper ───────────────────────────────────────────────────────────────

def _docx_to_text(source) -> str:
    """Extract plain text from a .docx file or bytes.

    Iterates paragraphs first, then tables, so structured CVs that use tables
    for employment history are fully captured.

    Args:
        source: A file path (str / os.PathLike) or raw ``bytes``.

    Returns:
        Extracted plain text with one entry per paragraph / table row.

    Raises:
        ImportError: If ``python-docx`` is not installed.
    """
    try:
        import docx as _docx
    except ImportError:
        raise ImportError(
            "python-docx is required to read .docx files.\n"
            "Install it with:  pip install python-docx\n"
            "              or:  pip install \"resume-rdf[docx]\""
        )

    if isinstance(source, (str, os.PathLike)):
        doc = _docx.Document(str(source))
    else:
        doc = _docx.Document(io.BytesIO(source))

    lines: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append("  |  ".join(cells))

    return "\n".join(lines)


# ── content builders ──────────────────────────────────────────────────────────

def build_user_content_from_path(file_path: str, extra_context: str) -> list:
    """Build the ``messages[0].content`` block from a file path.

    Supports PDF (sent as a base64 document block), Word documents (text
    extracted via python-docx), and plain text / Markdown (sent as a text
    block).

    Args:
        file_path: Absolute or relative path to a ``.pdf``, ``.docx``,
            ``.txt``, or ``.md`` file.
        extra_context: Optional free-text hint appended to the user message
            (e.g. preferred output language, main industry sectors).

    Returns:
        A list of content blocks ready for the Anthropic ``messages`` API.

    Raises:
        ImportError: If a ``.docx`` file is provided but python-docx is not
            installed.
    """
    ext = os.path.splitext(file_path)[1].lower()
    suffix = (
        f"\n\nAdditional context from the CV owner: {extra_context}"
        if extra_context.strip()
        else ""
    )
    if ext == ".pdf":
        with open(file_path, "rb") as f:
            b64 = base64.standard_b64encode(f.read()).decode("utf-8")
        return [
            {
                "type": "document",
                "source": {
                    "type": "base64",
                    "media_type": "application/pdf",
                    "data": b64,
                },
            },
            {
                "type": "text",
                "text": f"Parse this CV into Turtle RDF exactly as instructed.{suffix}",
            },
        ]
    if ext == ".docx":
        text = _docx_to_text(file_path)
    else:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
    return [
        {
            "type": "text",
            "text": (
                f"Parse this CV into Turtle RDF exactly as instructed."
                f"{suffix}\n\nCV content:\n\n{text}"
            ),
        }
    ]


def build_user_content_from_bytes(
    file_bytes: bytes, file_name: str, extra_context: str
) -> list:
    """Build the ``messages[0].content`` block from raw bytes.

    Suitable for Streamlit or other contexts where the file is already in
    memory.

    Args:
        file_bytes: Raw file contents.
        file_name: Original filename, used to determine the media type
            (``*.pdf`` → document block; ``*.docx`` → text extracted via
            python-docx; anything else → UTF-8 text block).
        extra_context: Optional free-text hint appended to the user message.

    Returns:
        A list of content blocks ready for the Anthropic ``messages`` API.

    Raises:
        ImportError: If a ``.docx`` file is provided but python-docx is not
            installed.
    """
    suffix = (
        f"\n\nAdditional context from the CV owner: {extra_context}"
        if extra_context.strip()
        else ""
    )
    ext = file_name.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        b64 = base64.standard_b64encode(file_bytes).decode("utf-8")
        return [
            {
                "type": "document",
                "source": {
                    "type": "base64",
                    "media_type": "application/pdf",
                    "data": b64,
                },
            },
            {
                "type": "text",
                "text": f"Parse this CV into Turtle RDF exactly as instructed.{suffix}",
            },
        ]
    if ext == "docx":
        text = _docx_to_text(file_bytes)
    else:
        text = file_bytes.decode("utf-8", errors="replace")
    return [
        {
            "type": "text",
            "text": (
                f"Parse this CV into Turtle RDF exactly as instructed."
                f"{suffix}\n\nCV content:\n\n{text}"
            ),
        }
    ]


def strip_fences(text: str) -> str:
    """Remove Markdown code fences that the model may have included despite instructions.

    Args:
        text: Raw model output.

    Returns:
        Turtle text with leading/trailing fence lines removed.
    """
    lines = text.strip().splitlines()
    if lines and lines[0].startswith("```"):
        lines = lines[1:]
    if lines and lines[-1].startswith("```"):
        lines = lines[:-1]
    return "\n".join(lines).strip()


def count_triples(ttl: str) -> int:
    """Estimate the number of RDF statements in a Turtle document.

    Counts lines that end with ``.`` or ``;`` and are not comments or prefix
    declarations.  This is a heuristic — it under-counts multi-value object
    lists and over-counts nothing — but it is fast and good enough for display.

    Args:
        ttl: Turtle RDF text.

    Returns:
        Approximate number of triple statements.
    """
    return sum(
        1
        for line in ttl.splitlines()
        if line.strip()
        and not line.strip().startswith(("#", "@"))
        and (line.rstrip().endswith(".") or line.rstrip().endswith(";"))
    )


def extract_person_name(ttl: str) -> str | None:
    """Extract the ``foaf:name`` value from a Turtle document.

    Args:
        ttl: Turtle RDF text.

    Returns:
        The person's name string, or ``None`` if not found.
    """
    m = re.search(r'foaf:name\s+"([^"]+)"', ttl)
    return m.group(1) if m else None


def validate_turtle(ttl: str) -> bool:
    """Parse *ttl* with rdflib to check for syntax errors.

    Requires ``pip install rdflib``.  If rdflib is not installed the function
    prints a notice and returns ``True`` (non-blocking).

    Args:
        ttl: Turtle RDF text to validate.

    Returns:
        ``True`` if the document is valid (or rdflib is unavailable),
        ``False`` on a parse error.
    """
    try:
        from rdflib import Graph  # type: ignore
        g = Graph()
        g.parse(data=ttl, format="turtle")
        print(f"  v  rdflib: {len(g)} triples parsed successfully.")
        return True
    except ImportError:
        print("  .  rdflib not installed — skipping validation.")
        print("     Install with:  pip install rdflib")
        return True
    except Exception as exc:
        import sys
        print(f"  x  rdflib validation error: {exc}", file=sys.stderr)
        return False
